from __future__ import annotations

import copy
import hashlib
import json
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import plotly.graph_objects as go
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from langchain_core.messages import HumanMessage, SystemMessage

from backend.agent.llm_client import make_reasoning_llm
from backend.core.config import settings
from backend.tools.impl.plotly_tool import CHART_COLORWAY, _plotly_sequence, apply_default_chart_style

_ARTIFACT_RE = re.compile(
    r"\[\s*artifact\s*:\s*([A-Za-z0-9_\-]+)\s*\]",
    re.IGNORECASE,
)
_BOLD_HEADING_RE = re.compile(r"^\s*\*\*(.+?)\*\*\s*$")
_INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

# Reuse the same palette as the UI (`plotly_tool.CHART_COLORWAY`).
_PLOTLY_COLORWAY: tuple[str, ...] = CHART_COLORWAY
_BOARD_BAR_COLOR = "#2563eb"
_NOTE_GENERIC_TITLE = "Аналитическая записка"


@dataclass
class ReportBuildResult:
    file_name: str
    file_path: str
    download_url: str


def _artifact_name_from_message_refs(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [item.strip() for item in value if isinstance(item, str) and item.strip()]


def _is_plotly_format(value: str) -> bool:
    normalized = str(value or "").strip().lower().replace("_", "-")
    return normalized == "plotly-json"


def _safe_file_stem(value: str, fallback: str = "chart") -> str:
    stem = re.sub(r"[^\w_.-]+", "_", str(value or ""), flags=re.UNICODE).strip("._-")
    return stem or fallback


def _build_artifact_index(artifacts: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    result: dict[str, dict[str, Any]] = {}
    for art in artifacts:
        if not isinstance(art, dict):
            continue

        art_type = str(art.get("type") or art.get("presentation_type") or "").strip().lower()
        title = str(art.get("text") or art.get("title") or "").strip()
        data = art.get("data") if isinstance(art.get("data"), dict) else {}
        data_format = str(data.get("format") or "").strip().lower()
        payload = data.get("data")
        meta = art.get("meta") if isinstance(art.get("meta"), dict) else {}

        artifact_name = ""
        if isinstance(meta.get("artifact_name"), str) and str(meta.get("artifact_name")).strip():
            artifact_name = str(meta["artifact_name"]).strip()
        elif title:
            artifact_name = title

        if not artifact_name:
            continue

        result[artifact_name] = {
            "artifact_name": artifact_name,
            "artifact_type": art_type or "unknown",
            "title": title or artifact_name,
            "data_format": data_format,
            "payload": payload,
        }
    return result


def _plot_context_block(artifact: dict[str, Any]) -> str:
    payload = artifact.get("payload")
    fig = None
    try:
        if isinstance(payload, dict):
            fig = go.Figure(payload)
    except Exception:
        fig = None

    chart_title = artifact.get("title") or artifact.get("artifact_name") or "-"
    x_title = "-"
    y_title = "-"
    trace_types = "unknown"
    series_names = "-"

    if fig is not None:
        try:
            layout = fig.layout
            chart_title = getattr(getattr(layout, "title", None), "text", "") or chart_title
            x_title = getattr(getattr(getattr(layout, "xaxis", None), "title", None), "text", "") or "-"
            y_title = getattr(getattr(getattr(layout, "yaxis", None), "title", None), "text", "") or "-"
            trace_types = ", ".join(sorted({getattr(trace, "type", "unknown") for trace in fig.data})) or "unknown"  # noqa: E501
            names = [str(getattr(trace, "name", "")).strip() for trace in fig.data]
            names = [x for x in names if x]
            if names:
                series_names = ", ".join(names)
        except Exception:
            pass

    return (
        f"- artifact_name: {artifact['artifact_name']}\n"
        f"  artifact_type: plot\n"
        f"  chart_title: {chart_title}\n"
        f"  x_title: {x_title}\n"
        f"  y_title: {y_title}\n"
        f"  trace_types: {trace_types}\n"
        f"  series_names: {series_names}"
    )


def _build_prompt(chat_history: list[dict[str, Any]], artifacts: list[dict[str, Any]]) -> str:
    artifact_index = _build_artifact_index(artifacts)

    lines: list[str] = []
    lines.append("[ИСТОРИЯ ЧАТА]")
    for idx, msg in enumerate(chat_history, start=1):
        if not isinstance(msg, dict):
            continue

        role = str(msg.get("role") or "").strip()
        content = str(msg.get("content") or "").strip()
        if not content:
            continue

        lines.append(f"{idx}. role={role}")
        lines.append(content)

        linked_names = _artifact_name_from_message_refs(msg.get("artifacts"))
        linked_plot_names = [name for name in linked_names if name in artifact_index]

        if linked_plot_names:
            lines.append("Связанные графики:")
            for name in linked_plot_names:
                art = artifact_index[name]
                if _is_plotly_format(str(art.get("data_format") or "")):
                    lines.append(_plot_context_block(art))
        lines.append("")

    lines.append("[ВСЕ ДОСТУПНЫЕ ГРАФИКИ]")
    for art in artifact_index.values():
        if not _is_plotly_format(str(art.get("data_format") or "")):
            continue
        lines.append(_plot_context_block(art))
        lines.append("")

    return "\n".join(lines).strip()


def _build_llm():
    return make_reasoning_llm(
        provider=settings.llm_provider,
        model=settings.llm_model,
        base_url=settings.llm_base_url,
        api_key=settings.llm_api_key,
        enable_thinking=False,
        temperature=getattr(settings, "llm_temperature_chat", 0.3),
        max_tokens=getattr(settings, "llm_max_tokens_default", 4096),
        streaming=False,
        timeout=max(60.0, float(getattr(settings, "agent_step_timeout_sec", 60.0))),
        top_p=1.0,
        top_k=getattr(settings, "llm_top_k", 0),
        num_ctx=getattr(settings, "llm_num_ctx", 0),
        presence_penalty=0.0,
        chat_template_kwargs_enabled=getattr(settings, "llm_chat_template_kwargs_enabled", False),
    )


_SYSTEM_PROMPT = """
Ты пишешь деловой отчет по истории чата и связанным графикам.

Правила:
1. Пиши отчет на русском языке.
2. Используй markdown-подобный формат.
3. Заголовки разделов выделяй строго так: **Текст заголовка**
4. Для вставки графиков используй только формат: [artifact:artifact_name]
5. Используй только реально доступные artifact_name из контекста.
6. Не придумывай новые артефакты.
7. Не вставляй таблицы.
8. Структура: **Название отчета**, краткое введение, основные выводы, рекомендации.
9. Возвращай только текст отчета без пояснений.
""".strip()


def _validate_artifact_refs(report_text: str, artifacts: list[dict[str, Any]]) -> None:
    index = _build_artifact_index(artifacts)
    allowed = set(index.keys())

    found = [name.strip() for name in _ARTIFACT_RE.findall(report_text)]
    invented = sorted({name for name in found if name not in allowed})

    if invented:
        raise ValueError(f"LLM invented artifact names: {invented}")


def _with_alpha(hex_color: str, alpha: float) -> str:
    normalized = str(hex_color or "").replace("#", "").strip()
    if len(normalized) == 3:
        normalized = "".join(ch * 2 for ch in normalized)

    try:
        value = int(normalized, 16)
    except ValueError:
        return f"rgba(148, 163, 184, {alpha})"

    r = (value >> 16) & 255
    g = (value >> 8) & 255
    b = value & 255
    return f"rgba({r}, {g}, {b}, {alpha})"


def _resolve_trace_color(name_raw: Any, index: int) -> str:
    name = str(name_raw or "").lower()

    if re.search(r"(anomaly|outlier|alert|аномал)", name):
        return "#dc2626"

    if re.search(r"(forecast|prediction|pred|yhat|plan|expected|прогноз|план)", name):
        return "#7c3aed"

    if re.search(r"(fact|actual|real|observed|факт)", name) or name == "y":
        return "#2563eb"

    if re.search(r"(lower|upper|bound|interval|confidence|band|ci)", name):
        return "#94a3b8"

    return _PLOTLY_COLORWAY[index % len(_PLOTLY_COLORWAY)]


def _normalize_plotly_trace(raw_trace: Any, index: int, *, is_dark: bool = False) -> dict[str, Any]:
    trace = dict(raw_trace or {}) if isinstance(raw_trace, dict) else {}

    name = str(trace.get("name") or "").lower()
    trace_type = str(trace.get("type") or "scatter").lower()
    mode = str(trace.get("mode") or "").lower()
    fill = str(trace.get("fill") or "").lower()

    color = _resolve_trace_color(trace.get("name"), index)

    is_band = (
        re.search(r"(lower|upper|bound|interval|confidence|band|ci)", name) is not None
        or fill in {"tonexty", "tozeroy"}
    )

    if trace_type in {"scatter", "scattergl", ""}:
        base_line = trace.get("line") if isinstance(trace.get("line"), dict) else {}
        trace["line"] = {
            **base_line,
            "color": color,
            "width": 1.6 if is_band else 2.4,
        }

        if "markers" in mode:
            base_marker = trace.get("marker") if isinstance(trace.get("marker"), dict) else {}
            trace["marker"] = {
                **base_marker,
                "color": color,
                "line": {"width": 0},
            }

        if is_band:
            trace["fillcolor"] = _with_alpha(color, 0.18 if is_dark else 0.12)

    elif trace_type == "pie":
        try:
            label_count = max(1, len(_plotly_sequence(trace.get("labels"))))
        except Exception:
            label_count = 1
        colors = [
            _PLOTLY_COLORWAY[(index + point_index) % len(_PLOTLY_COLORWAY)]
            for point_index in range(label_count)
        ]
        marker = trace.get("marker") if isinstance(trace.get("marker"), dict) else {}
        trace["marker"] = {
            **marker,
            "colors": colors,
            "line": {"color": "#ffffff", "width": 1},
        }

    elif trace_type in {"bar", "histogram"}:
        orientation = str(trace.get("orientation") or "").lower()
        primary = trace.get("y") if orientation != "h" else trace.get("x")
        try:
            point_count = max(1, len(_plotly_sequence(primary)))
        except Exception:
            point_count = 1
        colors = [
            _PLOTLY_COLORWAY[(index + point_index) % len(_PLOTLY_COLORWAY)]
            for point_index in range(point_count)
        ]
        base_marker = trace.get("marker") if isinstance(trace.get("marker"), dict) else {}
        trace["marker"] = {
            **base_marker,
            "color": colors if point_count > 1 else colors[0],
            "line": {"width": 0},
            "opacity": 0.9,
            "cornerradius": 6,
        }

    return trace


def _normalize_axis_layout(
    raw_axis: Any,
    *,
    text: str,
    muted: str,
    grid: str,
    zero: str,
) -> dict[str, Any]:
    axis = dict(raw_axis or {}) if isinstance(raw_axis, dict) else {}

    tickfont = axis.get("tickfont") if isinstance(axis.get("tickfont"), dict) else {}
    title = axis.get("title") if isinstance(axis.get("title"), dict) else {}
    title_font = title.get("font") if isinstance(title.get("font"), dict) else {}

    return {
        **axis,
        "automargin": True,
        "gridcolor": grid,
        "zerolinecolor": zero,
        "tickfont": {
            **tickfont,
            "color": muted,
        },
        "title": {
            **title,
            "font": {
                **title_font,
                "color": text,
            },
        },
    }


def _normalize_annotations(raw_annotations: Any, *, text: str) -> Any:
    if not isinstance(raw_annotations, list):
        return raw_annotations

    normalized: list[dict[str, Any]] = []
    for raw_ann in raw_annotations:
        if not isinstance(raw_ann, dict):
            continue

        font = raw_ann.get("font") if isinstance(raw_ann.get("font"), dict) else {}
        normalized.append({
            **raw_ann,
            "font": {
                **font,
                "color": text,
            },
        })

    return normalized


def _style_plotly_payload_for_report(payload: dict[str, Any]) -> dict[str, Any]:
    """
    Apply frontend-like Plotly styling before static PNG export.

    Frontend renders charts inside a responsive card, so it can remove fixed
    width/height. DOCX export needs deterministic image dimensions, therefore
    width/height are set here explicitly.
    """
    is_dark = False

    frame_bg = "#09090b" if is_dark else "#ffffff"
    plot_bg = frame_bg
    text = "#fafafa" if is_dark else "#18181b"
    muted = "#a1a1aa" if is_dark else "#717182"
    grid = "rgba(63,63,70,0.5)" if is_dark else "rgba(0,0,0,0.10)"
    zero = "#3f3f46" if is_dark else "rgba(0,0,0,0.18)"

    normalized = copy.deepcopy(payload)

    raw_data = normalized.get("data")
    traces = raw_data if isinstance(raw_data, list) else []
    normalized["data"] = [
        _normalize_plotly_trace(trace, index, is_dark=is_dark)
        for index, trace in enumerate(traces)
    ]

    base_layout = normalized.get("layout")
    base_layout = dict(base_layout or {}) if isinstance(base_layout, dict) else {}

    base_font = base_layout.get("font") if isinstance(base_layout.get("font"), dict) else {}

    base_title = base_layout.get("title") if isinstance(base_layout.get("title"), dict) else {}
    base_title_font = base_title.get("font") if isinstance(base_title.get("font"), dict) else {}

    base_legend = base_layout.get("legend") if isinstance(base_layout.get("legend"), dict) else {}
    base_legend_font = base_legend.get("font") if isinstance(base_legend.get("font"), dict) else {}

    base_legend_title = base_legend.get("title") if isinstance(base_legend.get("title"), dict) else {}
    base_legend_title_font = (
        base_legend_title.get("font")
        if isinstance(base_legend_title.get("font"), dict)
        else {}
    )

    base_margin = base_layout.get("margin") if isinstance(base_layout.get("margin"), dict) else {}

    normalized["layout"] = {
        **base_layout,
        "width": 1200,
        "height": 720,
        "autosize": False,
        "colorway": list(_PLOTLY_COLORWAY),
        "paper_bgcolor": frame_bg,
        "plot_bgcolor": plot_bg,
        "font": {
            **base_font,
            "color": text,
            "family": "ui-sans-serif, system-ui, sans-serif",
        },
        "title": {
            **base_title,
            "font": {
                **base_title_font,
                "color": text,
                "family": "ui-sans-serif, system-ui, sans-serif",
            },
        },
        "legend": {
            **base_legend,
            "bgcolor": "rgba(0,0,0,0)",
            "font": {
                **base_legend_font,
                "color": muted,
            },
            "title": {
                **base_legend_title,
                "font": {
                    **base_legend_title_font,
                    "color": muted,
                },
            },
        },
        "xaxis": _normalize_axis_layout(
            base_layout.get("xaxis"),
            text=text,
            muted=muted,
            grid=grid,
            zero=zero,
        ),
        "yaxis": _normalize_axis_layout(
            base_layout.get("yaxis"),
            text=text,
            muted=muted,
            grid=grid,
            zero=zero,
        ),
        "annotations": _normalize_annotations(
            base_layout.get("annotations"),
            text=text,
        ),
        "margin": {
            "l": 44,
            "r": 28,
            "t": 44,
            "b": 44,
            **base_margin,
        },
        "bargap": 0.32,
        "bargroupgap": 0.14,
    }

    return normalized


def _apply_board_bar_colors(fig: go.Figure) -> None:
    """Board/PDF export: one color per bar chart (UI uses the same for structure bars)."""
    for trace in fig.data:
        if str(getattr(trace, "type", "") or "").lower() != "bar":
            continue
        marker = getattr(trace, "marker", None)
        marker_dict = dict(marker) if isinstance(marker, dict) else {}
        marker_dict["color"] = _BOARD_BAR_COLOR
        marker_dict.setdefault("line", {"width": 0})
        trace.marker = marker_dict  # type: ignore[attr-defined]


def _render_plot_png(
    artifact: dict[str, Any],
    export_dir: Path,
    *,
    board_export: bool = False,
) -> Path:
    payload = artifact.get("payload")
    if not isinstance(payload, dict):
        raise ValueError(f"Artifact {artifact['artifact_name']} has no plotly-json payload")

    fig = apply_default_chart_style(go.Figure(copy.deepcopy(payload)))
    if board_export:
        _apply_board_bar_colors(fig)
    fig.update_layout(
        width=1200,
        height=720,
        autosize=False,
        paper_bgcolor="#ffffff",
        plot_bgcolor="#ffffff",
    )

    safe_name = _safe_file_stem(str(artifact.get("artifact_name") or "chart"))
    digest_source = {
        "artifact_name": str(artifact.get("artifact_name") or ""),
        "title": str(artifact.get("title") or ""),
        "payload": payload,
    }
    digest = hashlib.sha1(
        json.dumps(digest_source, ensure_ascii=False, sort_keys=True, default=str).encode("utf-8")
    ).hexdigest()[:10]
    out_path = export_dir / f"{safe_name}_{digest}.png"

    fig.write_image(
        str(out_path),
        format="png",
        width=1200,
        height=720,
        scale=2,
    )
    return out_path


_HEADING_MD_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_BULLET_MD_RE = re.compile(r"^[-*•]\s+(.+)$")
_NUMBERED_MD_RE = re.compile(r"^\d+[.)]\s+(.+)$")
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_LINK_MD_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")


def _clean_inline_markdown(text: str) -> str:
    cleaned = str(text or "").strip()
    cleaned = _LINK_MD_RE.sub(r"\1", cleaned)
    cleaned = _INLINE_CODE_RE.sub(r"\1", cleaned)
    return cleaned.strip()


def _strip_markdown_syntax(text: str) -> str:
    cleaned = _clean_inline_markdown(text)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.+?)__", r"\1", cleaned)
    return re.sub(r"\s{2,}", " ", cleaned).strip()


def _add_runs_with_bold(paragraph: Any, text: str, *, base_size: int = 11) -> None:
    pos = 0
    for match in _INLINE_BOLD_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.size = Pt(base_size)
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        bold_run.font.size = Pt(base_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(base_size)


def _add_board_cover(doc: Document, title: str) -> None:
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(_strip_markdown_syntax(title) or "Отчёт по визуализациям")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stamp = datetime.now(timezone.utc).astimezone().strftime("%d.%m.%Y %H:%M")
    sub_run = subtitle.add_run(f"Сформировано: {stamp}")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x71, 0x71, 0x82)
    doc.add_paragraph("")


def _add_note_content_to_docx(doc: Document, content: str) -> None:
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        heading_match = _HEADING_MD_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            p = doc.add_paragraph()
            run = p.add_run(_strip_markdown_syntax(heading_match.group(2)))
            run.bold = True
            run.font.size = Pt(16 if level <= 2 else 13)
            run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)
            continue

        heading_bold = _BOLD_HEADING_RE.match(stripped)
        if heading_bold:
            p = doc.add_paragraph()
            run = p.add_run(_strip_markdown_syntax(heading_bold.group(1)))
            run.bold = True
            run.font.size = Pt(14)
            continue

        bullet_match = _BULLET_MD_RE.match(stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, _clean_inline_markdown(bullet_match.group(1)))
            continue

        numbered_match = _NUMBERED_MD_RE.match(stripped)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, _clean_inline_markdown(numbered_match.group(1)))
            continue

        p = doc.add_paragraph()
        _add_runs_with_bold(p, _clean_inline_markdown(stripped))


def _add_markdownish_paragraph(doc: Document, line: str) -> None:
    _add_note_content_to_docx(doc, str(line or ""))


def _render_docx(report_text: str, artifacts: list[dict[str, Any]], out_path: Path) -> None:
    artifact_index = _build_artifact_index(artifacts)
    export_dir = out_path.parent / f"{out_path.stem}_assets"
    export_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    for raw_line in report_text.splitlines():
        line = raw_line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        match = _ARTIFACT_RE.fullmatch(line)
        if match:
            artifact_name = match.group(1).strip()
            artifact = artifact_index.get(artifact_name)

            if artifact is None:
                p = doc.add_paragraph()
                run = p.add_run(f"[График не найден: {artifact_name}]")
                run.italic = True
                run.font.size = Pt(10)
                continue

            if not _is_plotly_format(str(artifact.get("data_format") or "")):
                p = doc.add_paragraph()
                run = p.add_run(f"[Артефакт не является графиком: {artifact_name}]")
                run.italic = True
                run.font.size = Pt(10)
                continue

            png_path = _render_plot_png(artifact, export_dir)
            doc.add_picture(str(png_path), width=Inches(6.5))
            continue

        _add_markdownish_paragraph(doc, line)

    doc.save(str(out_path))


def build_report_docx(
    *,
    session_id: str,
    chat_history: list[dict[str, Any]],
    artifacts: list[dict[str, Any]],
    output_dir: Path,
    base_download_url: str,
) -> ReportBuildResult:
    llm = _build_llm()
    prompt = _build_prompt(chat_history, artifacts)

    response = llm.invoke([
        SystemMessage(content=_SYSTEM_PROMPT),
        HumanMessage(content=prompt),
    ])
    report_text = str(getattr(response, "content", "") or "").strip()
    if not report_text:
        raise ValueError("LLM returned empty report")

    _validate_artifact_refs(report_text, artifacts)

    file_name = f"report_{session_id}_{uuid.uuid4().hex[:8]}.docx"
    file_path = output_dir / file_name
    output_dir.mkdir(parents=True, exist_ok=True)

    _render_docx(report_text, artifacts, file_path)

    return ReportBuildResult(
        file_name=file_name,
        file_path=str(file_path),
        download_url=f"{base_download_url.rstrip('/')}/{file_name}",
    )


def _note_markdown_from_artifact(artifact: dict[str, Any]) -> str:
    data = artifact.get("data") if isinstance(artifact.get("data"), dict) else {}
    if str(data.get("format") or "").strip().lower() == "markdown":
        payload = data.get("data")
        if isinstance(payload, dict):
            return str(payload.get("content") or "").strip()
    return str(artifact.get("text") or "").strip()


def _normalize_heading_key(text: str) -> str:
    cleaned = _strip_markdown_syntax(text).strip().lower()
    cleaned = re.sub(r"^\d+[.)]\s*", "", cleaned).strip()
    return cleaned


def _note_export_heading(artifact: dict[str, Any]) -> str:
    meta = artifact.get("meta") if isinstance(artifact.get("meta"), dict) else {}
    user_question = str(meta.get("user_question") or "").strip()
    if user_question:
        return user_question
    card_title = _strip_markdown_syntax(str(artifact.get("text") or "").strip())
    if card_title and _normalize_heading_key(card_title) not in {"суть", "ключевые цифры", "инсайты"}:
        return card_title
    return _NOTE_GENERIC_TITLE


def _strip_duplicate_note_heading(content: str, heading: str) -> str:
    """Remove leading title lines / prefix duplicated in export heading (e.g. «Суть Суть»)."""
    text = str(content or "").strip()
    if not text:
        return text

    heading_key = _normalize_heading_key(heading)
    generic_keys = {
        _normalize_heading_key(_NOTE_GENERIC_TITLE),
        "суть",
        "ключевые цифры",
        "инсайты",
        "выводы",
        "графики и артефакты",
    }
    if not heading_key:
        heading_key = "суть"

    lines = text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)

    changed = True
    while lines and changed:
        changed = False
        first = lines[0].strip()
        first_key = _normalize_heading_key(first)
        if first_key and first_key == heading_key:
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
            changed = True
            continue
        if _BOLD_HEADING_RE.match(first) and first_key in generic_keys:
            lines.pop(0)
            while lines and not lines[0].strip():
                lines.pop(0)
            changed = True

    text = "\n".join(lines).strip()
    if not text:
        return ""

    plain_heading = _strip_markdown_syntax(heading).strip()
    if plain_heading:
        lowered = text.lower()
        prefix = plain_heading.lower()
        if lowered.startswith(prefix):
            rest = text[len(plain_heading) :].lstrip(" \t:—-\n")
            if rest:
                text = rest

    return text.strip()


def _board_export_sections(
    artifacts: list[dict[str, Any]],
    sections: list[dict[str, Any]] | None,
) -> list[tuple[str, list[dict[str, Any]]]]:
    if not sections:
        return [("", artifacts)]

    by_id: dict[str, dict[str, Any]] = {}
    for artifact in artifacts:
        if not isinstance(artifact, dict):
            continue
        art_id = str(artifact.get("id") or "").strip()
        if art_id:
            by_id[art_id] = artifact

    used_ids: set[str] = set()
    grouped: list[tuple[str, list[dict[str, Any]]]] = []

    for raw_section in sections:
        if not isinstance(raw_section, dict):
            continue
        label = str(raw_section.get("label") or "").strip()
        section_items: list[dict[str, Any]] = []
        for raw_id in raw_section.get("artifact_ids") or []:
            art_id = str(raw_id or "").strip()
            if not art_id or art_id in used_ids:
                continue
            artifact = by_id.get(art_id)
            if artifact is None:
                continue
            used_ids.add(art_id)
            section_items.append(artifact)
        if label or section_items:
            grouped.append((label, section_items))

    remainder = [
        artifact
        for artifact in artifacts
        if isinstance(artifact, dict)
        and str(artifact.get("id") or "").strip() not in used_ids
    ]
    if remainder:
        grouped.append(("", remainder))

    return grouped or [("", artifacts)]


def _artifact_plot_entry(
    artifact: dict[str, Any],
    artifact_index: dict[str, dict[str, Any]],
) -> dict[str, Any] | None:
    data = artifact.get("data") if isinstance(artifact.get("data"), dict) else {}
    data_format = str(data.get("format") or "").strip().lower()
    payload = data.get("data")
    if _is_plotly_format(data_format) and isinstance(payload, dict):
        meta = artifact.get("meta") if isinstance(artifact.get("meta"), dict) else {}
        title = str(artifact.get("text") or "").strip()
        artifact_name = str(meta.get("artifact_name") or artifact.get("id") or title or "chart").strip()
        return {
            "artifact_name": artifact_name,
            "artifact_type": "plot",
            "title": title or artifact_name,
            "data_format": data_format,
            "payload": payload,
        }

    meta = artifact.get("meta") if isinstance(artifact.get("meta"), dict) else {}
    candidates = [
        str(meta.get("artifact_name") or "").strip(),
        str(artifact.get("text") or "").strip(),
        str(artifact.get("id") or "").strip(),
    ]
    for name in candidates:
        if name and name in artifact_index:
            return artifact_index[name]
    for entry in artifact_index.values():
        if _is_plotly_format(str(entry.get("data_format") or "")):
            title = str(artifact.get("text") or "").strip()
            if title and title == str(entry.get("title") or entry.get("artifact_name")):
                return entry
    return None


def _add_board_question_heading_docx(doc: Document, label: str) -> None:
    cleaned = _strip_markdown_syntax(label).strip()
    if not cleaned:
        return
    heading = doc.add_paragraph()
    heading_run = heading.add_run(cleaned)
    heading_run.bold = True
    heading_run.font.size = Pt(16)
    heading_run.font.color.rgb = RGBColor(0x18, 0x18, 0x1B)
    doc.add_paragraph("")


def _append_board_artifact_docx(
    doc: Document,
    artifact: dict[str, Any],
    *,
    artifact_index: dict[str, dict[str, Any]],
    export_dir: Path,
) -> None:
    art_type = str(artifact.get("type") or "").strip().lower()
    section_title = str(artifact.get("text") or art_type or "Раздел").strip()

    if art_type == "plot":
        plot_entry = _artifact_plot_entry(artifact, artifact_index)
        if plot_entry is None:
            p = doc.add_paragraph()
            run = p.add_run(f"[График недоступен: {section_title}]")
            run.italic = True
            return
        chart_title = _strip_markdown_syntax(
            str(plot_entry.get("title") or section_title)
        )
        heading = doc.add_paragraph()
        heading_run = heading.add_run(chart_title)
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        heading_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        png_path = _render_plot_png(plot_entry, export_dir, board_export=True)
        doc.add_picture(str(png_path), width=Inches(6.5))
        caption = doc.add_paragraph()
        caption_run = caption.add_run(chart_title)
        caption_run.italic = True
        caption_run.font.size = Pt(9)
        caption_run.font.color.rgb = RGBColor(0x71, 0x71, 0x82)
        doc.add_paragraph("")
        return

    if art_type == "note":
        raw_content = _note_markdown_from_artifact(artifact)
        if not raw_content:
            return
        note_heading = _note_export_heading(artifact)
        content = _strip_duplicate_note_heading(raw_content, note_heading)
        if not content:
            return
        heading = doc.add_paragraph()
        heading_run = heading.add_run(_strip_markdown_syntax(note_heading))
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        _add_note_content_to_docx(doc, content)
        doc.add_paragraph("")
        return

    if art_type == "table":
        p = doc.add_paragraph()
        run = p.add_run(f"[Таблица: {section_title} — откройте в интерфейсе для деталей]")
        run.italic = True


def _render_board_docx(
    *,
    title: str,
    artifacts: list[dict[str, Any]],
    out_path: Path,
    sections: list[dict[str, Any]] | None = None,
) -> None:
    artifact_index = _build_artifact_index(artifacts)
    export_dir = out_path.parent / f"{out_path.stem}_assets"
    export_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    _add_board_cover(doc, title)

    for section_label, section_artifacts in _board_export_sections(artifacts, sections):
        if section_label:
            _add_board_question_heading_docx(doc, section_label)
        for artifact in section_artifacts:
            _append_board_artifact_docx(
                doc,
                artifact,
                artifact_index=artifact_index,
                export_dir=export_dir,
            )

    doc.save(str(out_path))


_PDF_FONT_NAME = "BoardExportDejaVu"


def _ensure_pdf_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    if _PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return _PDF_FONT_NAME

    candidates = (
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
    )
    for font_path in candidates:
        if font_path.exists():
            pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, str(font_path)))
            return _PDF_FONT_NAME
    raise RuntimeError("Не найден шрифт с поддержкой кириллицы для PDF-экспорта")


def _markdown_line_to_reportlab(line: str) -> str:
    from xml.sax.saxutils import escape

    stripped = str(line or "").strip()
    heading_match = _HEADING_MD_RE.match(stripped)
    if heading_match:
        stripped = heading_match.group(2)
    bullet_match = _BULLET_MD_RE.match(stripped)
    if bullet_match:
        stripped = f"• {bullet_match.group(1)}"
    numbered_match = _NUMBERED_MD_RE.match(stripped)
    if numbered_match:
        stripped = numbered_match.group(1)
    cleaned = _strip_markdown_syntax(stripped)
    escaped = escape(cleaned)
    return re.sub(
        r"\*\*(.+?)\*\*",
        lambda match: f"<b>{match.group(1)}</b>",
        escaped,
    )


def _render_board_pdf(
    *,
    title: str,
    artifacts: list[dict[str, Any]],
    out_path: Path,
    sections: list[dict[str, Any]] | None = None,
) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Image as PdfImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = _ensure_pdf_font()
    artifact_index = _build_artifact_index(artifacts)
    export_dir = out_path.parent / f"{out_path.stem}_assets"
    export_dir.mkdir(parents=True, exist_ok=True)

    title_style = ParagraphStyle(
        "BoardTitle",
        fontName=font_name,
        fontSize=18,
        leading=22,
        spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "BoardSection",
        fontName=font_name,
        fontSize=13,
        leading=17,
        spaceBefore=10,
        spaceAfter=6,
    )
    question_style = ParagraphStyle(
        "BoardQuestion",
        fontName=font_name,
        fontSize=15,
        leading=19,
        spaceBefore=14,
        spaceAfter=8,
        textColor="#18181b",
    )
    body_style = ParagraphStyle(
        "BoardBody",
        fontName=font_name,
        fontSize=11,
        leading=15,
        spaceAfter=4,
    )
    table_header_style = ParagraphStyle(
        "BoardTableHeader",
        fontName=font_name,
        fontSize=6.7,
        leading=8.2,
        textColor="#111827",
        wordWrap="CJK",
    )
    table_cell_style = ParagraphStyle(
        "BoardTableCell",
        fontName=font_name,
        fontSize=6.2,
        leading=7.6,
        textColor="#111827",
        wordWrap="CJK",
    )

    stamp = datetime.now(timezone.utc).astimezone().strftime("%d.%m.%Y %H:%M")
    subtitle_style = ParagraphStyle(
        "BoardSubtitle",
        fontName=font_name,
        fontSize=9,
        leading=12,
        textColor="#717182",
        spaceAfter=10,
    )
    story: list[Any] = [
        Paragraph(
            _markdown_line_to_reportlab(title.strip() or "Отчёт по визуализациям"),
            title_style,
        ),
        Paragraph(f"Сформировано: {stamp}", subtitle_style),
        Spacer(1, 0.4 * cm),
    ]

    for section_label, section_artifacts in _board_export_sections(artifacts, sections):
        if section_label:
            story.append(
                Paragraph(_markdown_line_to_reportlab(section_label), question_style)
            )
            story.append(Spacer(1, 0.2 * cm))

        for artifact in section_artifacts:
            if not isinstance(artifact, dict):
                continue
            art_type = str(artifact.get("type") or "").strip().lower()
            section_title = str(artifact.get("text") or art_type or "Раздел").strip()

            if art_type == "plot":
                plot_entry = _artifact_plot_entry(artifact, artifact_index)
                if plot_entry is None:
                    story.append(
                        Paragraph(
                            _markdown_line_to_reportlab(
                                f"[График недоступен: {section_title}]"
                            ),
                            body_style,
                        )
                    )
                    continue
                chart_title = str(plot_entry.get("title") or section_title)
                story.append(
                    Paragraph(_markdown_line_to_reportlab(chart_title), section_style)
                )
                png_path = _render_plot_png(plot_entry, export_dir, board_export=True)
                story.append(PdfImage(str(png_path), width=16 * cm, height=9 * cm))
                story.append(Spacer(1, 0.5 * cm))
                continue

            if art_type == "note":
                raw_content = _note_markdown_from_artifact(artifact)
                if not raw_content:
                    continue
                note_heading = _note_export_heading(artifact)
                content = _strip_duplicate_note_heading(raw_content, note_heading)
                if not content:
                    continue
                story.append(
                    Paragraph(_markdown_line_to_reportlab(note_heading), section_style)
                )
                for raw_line in content.splitlines():
                    line = raw_line.strip()
                    if not line:
                        story.append(Spacer(1, 0.15 * cm))
                        continue
                    heading_match = _HEADING_MD_RE.match(line)
                    line_style = (
                        section_style
                        if heading_match or _BOLD_HEADING_RE.match(line)
                        else body_style
                    )
                    story.append(Paragraph(_markdown_line_to_reportlab(line), line_style))
                story.append(Spacer(1, 0.35 * cm))
                continue

            if art_type == "table":
                table_payload = _split_table_payload(artifact)
                if table_payload is None:
                    continue
                columns, rows = table_payload
                if not columns:
                    continue
                table_title = _strip_markdown_syntax(section_title)
                story.append(
                    Paragraph(_markdown_line_to_reportlab(table_title), section_style)
                )
                max_cols = 12
                max_rows = 35
                display_columns = columns[:max_cols]
                display_rows = rows[:max_rows]
                table_data = [
                    [
                        Paragraph(
                            _markdown_line_to_reportlab(str(column))[:180],
                            table_header_style,
                        )
                        for column in display_columns
                    ]
                ]
                for row in display_rows:
                    table_data.append([
                        Paragraph(
                            _markdown_line_to_reportlab(str(value))[:220] if value is not None else "",
                            table_cell_style,
                        )
                        for value in row[:max_cols]
                    ])
                col_width = (26.7 * cm) / max(1, len(display_columns))
                pdf_table = Table(table_data, repeatRows=1, colWidths=[col_width] * len(display_columns))
                pdf_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#111827")),
                            ("FONTNAME", (0, 0), (-1, -1), font_name),
                            ("LEFTPADDING", (0, 0), (-1, -1), 2.5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 2.5),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                        ]
                    )
                )
                story.append(pdf_table)
                if len(rows) > max_rows or len(columns) > max_cols:
                    story.append(
                        Paragraph(
                            _markdown_line_to_reportlab(
                                f"Показан фрагмент: {min(len(rows), max_rows)} из {len(rows)} строк, "
                                f"{min(len(columns), max_cols)} из {len(columns)} столбцов. Полная таблица доступна в Excel."
                            ),
                            subtitle_style,
                        )
                    )
                story.append(Spacer(1, 0.45 * cm))

    document = SimpleDocTemplate(
        str(out_path),
        pagesize=landscape(A4),
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )
    document.build(story)


def _safe_excel_sheet_title(value: str, fallback: str, used: set[str]) -> str:
    cleaned = re.sub(r"[\[\]:*?/\\]+", " ", str(value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip() or fallback
    base = cleaned[:31].strip() or fallback[:31]
    candidate = base
    index = 2
    while candidate in used:
        suffix = f" {index}"
        candidate = f"{base[: 31 - len(suffix)]}{suffix}".strip()
        index += 1
    used.add(candidate)
    return candidate


def _split_table_payload(artifact: dict[str, Any]) -> tuple[list[str], list[list[Any]]] | None:
    data = artifact.get("data") if isinstance(artifact.get("data"), dict) else {}
    if str(data.get("format") or "").strip().lower() != "split":
        return None
    payload = data.get("export_data") or data.get("data")
    if not isinstance(payload, dict):
        return None
    raw_columns = payload.get("columns")
    raw_rows = payload.get("data")
    if not isinstance(raw_columns, list) or not isinstance(raw_rows, list):
        return None
    columns = [str(column) for column in raw_columns]
    rows: list[list[Any]] = []
    for raw_row in raw_rows:
        if isinstance(raw_row, list):
            rows.append(raw_row)
        else:
            rows.append([raw_row])
    return columns, rows


def _autosize_excel_columns(sheet: Any, *, max_width: int = 48) -> None:
    for column_cells in sheet.columns:
        letter = column_cells[0].column_letter
        width = 10
        for cell in column_cells:
            value = "" if cell.value is None else str(cell.value)
            width = max(width, min(max_width, len(value) + 2))
        sheet.column_dimensions[letter].width = width


def _append_plot_data_sheet(
    *,
    workbook: Any,
    artifact: dict[str, Any],
    artifact_index: dict[str, dict[str, Any]],
    used_titles: set[str],
    summary_sheet: Any,
) -> None:
    plot_entry = _artifact_plot_entry(artifact, artifact_index)
    title = str(artifact.get("text") or "График").strip() or "График"
    summary_sheet.append([title, "График", "См. лист с данными графика"])
    if plot_entry is None:
        return
    payload = plot_entry.get("payload")
    try:
        fig = go.Figure(payload) if isinstance(payload, dict) else None
    except Exception:
        fig = None
    if fig is None:
        return

    sheet = workbook.create_sheet(_safe_excel_sheet_title(title, "График", used_titles))
    sheet.append(["Серия", "X", "Y"])
    for trace in fig.data:
        name = str(getattr(trace, "name", "") or title)
        xs = _plotly_sequence(getattr(trace, "x", None))
        ys = _plotly_sequence(getattr(trace, "y", None))
        max_len = max(len(xs), len(ys))
        for idx in range(max_len):
            sheet.append([
                name,
                xs[idx] if idx < len(xs) else None,
                ys[idx] if idx < len(ys) else None,
            ])
    _autosize_excel_columns(sheet)


def _append_planfact_help_sheet(
    *,
    workbook: Any,
    used_titles: set[str],
) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill

    sheet = workbook.create_sheet(
        _safe_excel_sheet_title("Как проверить", "Проверка", used_titles),
        1,
    )
    sheet.append(["Проверка расчёта план-факта"])
    sheet["A1"].font = Font(bold=True, size=14)
    sheet["A1"].fill = PatternFill(fill_type="solid", fgColor="DBEAFE")
    instructions = [
        "1. Откройте лист «Контроль»: каждая строка соответствует строке исходного результата.",
        "2. «Проверка Excel» пересчитывается по строкам листа «Расчетная детализация».",
        "3. Статус OK означает, что результат и пересчёт Excel совпали с точностью до 0,01.",
        "4. Если листы первички доступны, используйте на них тот же номер строки результата.",
        "В книгу включены только строки, участвовавшие в выгруженном результате.",
    ]
    for line in instructions:
        sheet.append([line])
    sheet.column_dimensions["A"].width = 110
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row):
        row[0].alignment = Alignment(wrap_text=True, vertical="top")


def _append_planfact_data_sheet(
    *,
    workbook: Any,
    title: str,
    payload: dict[str, Any],
    used_titles: set[str],
) -> None:
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    columns = [str(column) for column in payload.get("columns") or []]
    rows = payload.get("rows") or []
    if not columns:
        return

    sheet = workbook.create_sheet(_safe_excel_sheet_title(title, title, used_titles))
    header_row = 1
    sheet.append(columns)
    for cell in sheet[header_row]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(fill_type="solid", fgColor="DBEAFE")
    for row in rows:
        sheet.append(list(row) if isinstance(row, list | tuple) else [row])

    data_end = max(header_row, header_row + len(rows))
    last_column = get_column_letter(len(columns))
    sheet.auto_filter.ref = f"A{header_row}:{last_column}{data_end}"
    sheet.freeze_panes = f"A{header_row + 1}"
    _autosize_excel_columns(sheet)


def _append_planfact_validation_sheets(
    *,
    workbook: Any,
    validation_tables: dict[str, dict[str, Any]],
    used_titles: set[str],
    summary_sheet: Any,
) -> None:
    if not validation_tables:
        return
    if "Контроль" in validation_tables:
        _append_planfact_help_sheet(workbook=workbook, used_titles=used_titles)
    for title, payload in validation_tables.items():
        _append_planfact_data_sheet(
            workbook=workbook,
            title=title,
            payload=payload,
            used_titles=used_titles,
        )
        summary_sheet.append([title, "Проверка", f"Строк: {len(payload.get('rows') or [])}"])


def _render_board_xlsx(
    *,
    title: str,
    artifacts: list[dict[str, Any]],
    out_path: Path,
    sections: list[dict[str, Any]] | None = None,
    planfact_validation_tables: dict[str, dict[str, Any]] | None = None,
) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    artifact_index = _build_artifact_index(artifacts)
    workbook = Workbook()
    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    workbook.calculation.calcMode = "auto"
    summary = workbook.active
    summary.title = "Отчет"
    summary.append([_strip_markdown_syntax(title) or "Отчет"])
    summary.append([f"Сформировано: {datetime.now(timezone.utc).astimezone().strftime('%d.%m.%Y %H:%M')}"])
    summary.append([])
    summary.append(["Раздел", "Тип", "Комментарий"])
    for cell in summary[4]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(fill_type="solid", fgColor="E0F2FE")

    used_titles = {summary.title}
    _append_planfact_validation_sheets(
        workbook=workbook,
        validation_tables=planfact_validation_tables or {},
        used_titles=used_titles,
        summary_sheet=summary,
    )
    for section_label, section_artifacts in _board_export_sections(artifacts, sections):
        if section_label:
            summary.append([_strip_markdown_syntax(section_label), "", ""])

        for artifact in section_artifacts:
            if not isinstance(artifact, dict):
                continue
            art_type = str(artifact.get("type") or "").strip().lower()
            artifact_title = str(artifact.get("text") or art_type or "Раздел").strip()
            if art_type == "table":
                table_payload = _split_table_payload(artifact)
                if table_payload is None:
                    summary.append([artifact_title, "Таблица", "Не удалось прочитать split payload"])
                    continue
                columns, rows = table_payload
                if not columns or not rows:
                    summary.append([artifact_title, "Таблица", "Пропущена: нет данных"])
                    continue
                sheet = workbook.create_sheet(_safe_excel_sheet_title(artifact_title, "Таблица", used_titles))
                sheet.append(columns)
                for cell in sheet[1]:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(fill_type="solid", fgColor="DBEAFE")
                for row in rows:
                    sheet.append(row)
                sheet.freeze_panes = "A2"
                _autosize_excel_columns(sheet)
                summary.append([artifact_title, "Таблица", f"Строк: {len(rows)}"])
                continue
            if art_type == "plot":
                summary.append([artifact_title, "График", "Пропущен в Excel; графики доступны в PDF и дашборде"])
                continue
            if art_type == "note":
                content = _note_markdown_from_artifact(artifact)
                if content:
                    sheet = workbook.create_sheet(_safe_excel_sheet_title(artifact_title, "Заметка", used_titles))
                    sheet.append(["Текст"])
                    for line in content.splitlines():
                        sheet.append([_strip_markdown_syntax(line)])
                    _autosize_excel_columns(sheet)
                    summary.append([artifact_title, "Заметка", "См. отдельный лист"])

    _autosize_excel_columns(summary)
    workbook.save(str(out_path))


def build_board_export(
    *,
    title: str,
    artifacts: list[dict[str, Any]],
    output_dir: Path,
    export_format: str,
    sections: list[dict[str, Any]] | None = None,
    planfact_validation_tables: dict[str, dict[str, Any]] | None = None,
) -> ReportBuildResult:
    if not artifacts:
        raise ValueError("Нет артефактов для экспорта")

    normalized_format = str(export_format or "docx").strip().lower()
    if normalized_format not in {"docx", "pdf", "xlsx"}:
        raise ValueError(f"Unsupported export format: {export_format}")

    output_dir.mkdir(parents=True, exist_ok=True)
    extension = normalized_format
    file_name = f"board_report_{uuid.uuid4().hex[:8]}.{extension}"
    file_path = output_dir / file_name

    if normalized_format == "pdf":
        _render_board_pdf(
            title=title,
            artifacts=artifacts,
            out_path=file_path,
            sections=sections,
        )
    elif normalized_format == "xlsx":
        _render_board_xlsx(
            title=title,
            artifacts=artifacts,
            out_path=file_path,
            sections=sections,
            planfact_validation_tables=planfact_validation_tables,
        )
    else:
        _render_board_docx(
            title=title,
            artifacts=artifacts,
            out_path=file_path,
            sections=sections,
        )

    return ReportBuildResult(
        file_name=file_name,
        file_path=str(file_path),
        download_url="",
    )
